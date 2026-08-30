"""Mechanics for reworking Ayaan's CV without touching its formatting.

The CV is a .docx with its own fonts, spacing and margins. Never rebuild it in HTML —
edit the bullet text of a copy in place, then export through real Word so the PDF looks
exactly like the source document.

Usage (python is at %LOCALAPPDATA%\\Programs\\Python\\Python312\\python.exe):

    python tools/cv_tools.py inspect  <cv.docx>
    python tools/cv_tools.py export   <cv.docx> <out.pdf>

Or import it:

    import sys; sys.path.insert(0, "tools")
    import cv_tools as cv

    cv.inspect(SRC)                                  # see indices + char budgets
    cv.apply_bullets(SRC, DST, {14: "...", 15: "..."})
    cv.export_pdf(DST, PDF)                          # returns page count, verified
"""

import copy
import os
import shutil
import subprocess
import sys
import time

import docx


# ---------------------------------------------------------------- inspection

def paragraphs(path):
    return docx.Document(path).paragraphs


def inspect(path):
    """Print every paragraph with its index and character count.

    The index is what you pass to apply_bullets(); the count is the budget you must
    stay within for that bullet.
    """
    for i, p in enumerate(paragraphs(path)):
        text = p.text
        marker = "  <-- bullet" if _looks_like_bullet(p) else ""
        print(f"{i:>3}  [{len(text):>3}]  {text[:100]!r}{marker}")


def _looks_like_bullet(p):
    """Heuristic: bullets are body lines with no tab (headers use tabs for the
    right-aligned location/date column)."""
    return bool(p.text.strip()) and "\t" not in p.text and not p.text.isupper()


def bullet_budget(path, indices):
    """{index: original character count} — the ceiling for each replacement."""
    paras = paragraphs(path)
    return {i: len(paras[i].text) for i in indices}


# ------------------------------------------------------------------ editing

def _set_text(paragraph, text):
    """Replace a paragraph's text, keeping its exact formatting.

    Bullets are often split across several identically-formatted runs left over from
    past edits. Keeping run 0 and deleting the rest preserves font, size, bold and
    italic with no visible change.
    """
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


def apply_bullets(src, dst, replacements, drop=(), strict=True):
    """Copy src -> dst, reword the given paragraphs, optionally drop some.

    replacements: {paragraph_index: new_text}
    drop:         paragraph indices to delete entirely (applied after replacements)
    strict:       raise if any replacement exceeds its original character count

    Returns the list of paragraph indices that actually changed, so you can confirm
    nothing outside the intended set moved.
    """
    if strict:
        budget = bullet_budget(src, replacements)
        over = {i: (len(t), budget[i]) for i, t in replacements.items()
                if len(t) > budget[i]}
        if over:
            raise ValueError(
                "over budget (new, original): " + repr(over) +
                "\nThe CV fills exactly one page with no spare line: a bullet even ~16 "
                "characters longer can wrap onto an extra line and push the last line "
                "of the document onto page 2. Budget per bullet, not in aggregate.")

    shutil.copyfile(src, dst)
    document = docx.Document(dst)

    for index, text in replacements.items():
        _set_text(document.paragraphs[index], text)

    # Delete last, and from the bottom up, so earlier indices stay valid.
    for index in sorted(drop, reverse=True):
        element = document.paragraphs[index]._element
        element.getparent().remove(element)

    document.save(dst)

    before, after = paragraphs(src), paragraphs(dst)
    if drop:
        return sorted(set(replacements) | set(drop))
    changed = [i for i in range(len(before)) if before[i].text != after[i].text]
    return changed


def insert_after(dst, template_index, anchor_index, runs):
    """Clone an existing paragraph and insert it after another, then set its runs.

    Used to add a line that must match the look of an existing one (e.g. another
    "Label: value" line under ADDITIONAL INFORMATION). `runs` is {run_index: text};
    runs not listed keep their current text, so leave the tab run alone.
    """
    document = docx.Document(dst)
    clone = copy.deepcopy(document.paragraphs[template_index]._element)
    document.paragraphs[anchor_index]._element.addnext(clone)
    document.save(dst)

    document = docx.Document(dst)
    paragraph = document.paragraphs[anchor_index + 1]
    for run_index, text in runs.items():
        paragraph.runs[run_index].text = text
    document.save(dst)
    return anchor_index + 1


# ------------------------------------------------------------------- export

_PS_EXPORT = r"""
$ErrorActionPreference = 'Stop'
# Remember any Word already running so we only ever end the instance we start —
# killing all WINWORD would close a document the user has open.
$pre = @(Get-Process WINWORD -ErrorAction SilentlyContinue | Select-Object -ExpandProperty Id)
if (Test-Path '{out}') {{ Remove-Item '{out}' -Force }}
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
  $doc = $word.Documents.Open('{src}', [ref]$false, [ref]$true)
  Write-Output ("PAGES=" + $doc.ComputeStatistics(2))
  $doc.SaveAs([ref]'{out}', [ref]17)
  $doc.Close([ref]0)
}} finally {{
  # $word.Quit() routinely blocks for minutes under COM automation. The document is
  # saved and closed by this point, so end our own instance directly instead.
  Get-Process WINWORD -ErrorAction SilentlyContinue |
    Where-Object {{ $pre -notcontains $_.Id }} |
    Stop-Process -Force -ErrorAction SilentlyContinue
}}
"""


def export_pdf(docx_path, pdf_path, timeout=180):
    """Export through real Word (not a browser) and return the verified page count.

    Word's SaveAs can return before the file is flushed, so this waits for the size to
    stop changing before counting pages — a stale read once reported 1 page for a file
    that was really 2.
    """
    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path)
    script = _PS_EXPORT.format(src=docx_path.replace("'", "''"),
                               out=pdf_path.replace("'", "''"))
    proc = subprocess.run(["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
                          capture_output=True, text=True, timeout=timeout)
    if not os.path.isfile(pdf_path):
        raise RuntimeError(f"Word produced no PDF.\n{proc.stdout}\n{proc.stderr}")

    last = -1
    for _ in range(40):
        size = os.path.getsize(pdf_path)
        if size == last and size > 0:
            break
        last = size
        time.sleep(0.25)

    from pypdf import PdfReader
    return len(PdfReader(pdf_path).pages)


# ---------------------------------------------------------------------- cli

def _main(argv):
    if len(argv) >= 3 and argv[1] == "inspect":
        inspect(argv[2])
        return 0
    if len(argv) >= 4 and argv[1] == "export":
        pages = export_pdf(argv[2], argv[3])
        print(f"pages: {pages}")
        return 0 if pages == 1 else 1
    print(__doc__)
    return 2


if __name__ == "__main__":
    sys.exit(_main(sys.argv))
